"""
Document Analyzer
==================
Extracts text from uploaded documents (PDF, DOCX, MD, TXT)
and uses AI to identify requirements and functional flows.

TWO RESPONSIBILITIES:
1. extract_text() — Pull raw text from any supported file format
2. analyze() — Use AI to understand the document structure

SUPPORTED FORMATS:
- PDF  → PyMuPDF (fitz) — fast, accurate text extraction
- DOCX → python-docx — reads Word document paragraphs & tables
- MD   → direct file read (already text)
- TXT  → direct file read (already text)
"""

import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.ai.llm_client import llm_client

logger = logging.getLogger(__name__)

# Prompt for AI document analysis
DOCUMENT_ANALYSIS_PROMPT = """Analyze the following software document and extract structured information.

DOCUMENT TEXT:
{document_text}

Return a JSON object with these fields:
{{
    "sections": ["list of section titles/headings found in the document"],
    "requirements": ["list of functional and non-functional requirements extracted"],
    "functional_flows": ["list of user flows/scenarios described in the document"]
}}

Rules:
- Extract ALL requirements, even if they are implied
- Each requirement should be a clear, testable statement
- Functional flows should describe end-to-end user journeys
- If the document is short, still extract as much as possible

Return ONLY valid JSON, no extra text."""

ANALYSIS_SYSTEM_PROMPT = """You are an expert business analyst and QA engineer.
Your job is to read software requirement documents and extract structured information.
Always return valid JSON. Be thorough and precise."""


class DocumentAnalyzer:
    """Parses documents and extracts structured requirements using AI."""

    # ── Text Extraction (per file type) ────────────────────────────

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract raw text from a document file.

        Args:
            file_path: Absolute path to the uploaded file
            file_type: Extension without dot — "pdf", "docx", "md", "txt"

        Returns:
            Extracted text content as a single string
        """
        file_type = file_type.lower().lstrip(".")

        extractors = {
            "pdf": self._extract_from_pdf,
            "docx": self._extract_from_docx,
            "md": self._extract_from_text,
            "txt": self._extract_from_text,
        }

        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")

        text = await extractor(file_path)

        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return ""

        logger.info(
            f"Extracted {len(text)} characters from {Path(file_path).name}"
        )
        return text

    async def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using PyMuPDF (fitz).

        PyMuPDF reads each page and extracts text blocks.
        Handles multi-page documents, tables, and formatted text.
        """
        text_parts = []

        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        doc.close()

        return "\n\n".join(text_parts)

    async def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX using python-docx.

        Reads paragraphs and tables from Word documents.
        Preserves heading structure with markdown-style formatting.
        """
        doc = DocxDocument(file_path)
        text_parts = []

        # Extract paragraphs (including headings)
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Preserve heading hierarchy
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                text_parts.append(f"{'#' * level} {para.text}")
            else:
                text_parts.append(para.text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = [f"\n[Table {table_idx + 1}]"]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(cells))
            text_parts.append("\n".join(table_text))

        return "\n\n".join(text_parts)

    async def _extract_from_text(self, file_path: str) -> str:
        """Read text/markdown files directly — they're already text."""
        path = Path(file_path)
        return path.read_text(encoding="utf-8")

    # ── AI-Powered Analysis ────────────────────────────────────────

    async def analyze(self, text: str) -> dict:
        """
        Use AI to analyze extracted text and identify:
        - Document sections/headings
        - Functional & non-functional requirements
        - User flows and scenarios

        Args:
            text: Raw text extracted from the document

        Returns:
            {
                "sections": ["Login", "Registration", ...],
                "requirements": ["User must be able to...", ...],
                "functional_flows": ["User opens app → ...", ...]
            }
        """
        if not text.strip():
            return {
                "sections": [],
                "requirements": [],
                "functional_flows": [],
            }

        # Truncate very long documents to stay within LLM token limits
        max_chars = 30000  # ~7500 tokens — safe for most models
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[Document truncated for analysis]"

        prompt = DOCUMENT_ANALYSIS_PROMPT.format(document_text=text)

        try:
            response = await llm_client.generate(
                prompt=prompt,
                system_prompt=ANALYSIS_SYSTEM_PROMPT,
            )
            result = llm_client.parse_json_response(response)

            # Ensure all expected keys exist
            return {
                "sections": result.get("sections", []),
                "requirements": result.get("requirements", []),
                "functional_flows": result.get("functional_flows", []),
            }

        except Exception as e:
            logger.error(f"Document analysis failed: {e}")
            return {
                "sections": [],
                "requirements": [],
                "functional_flows": [],
                "error": str(e),
            }


# Single instance used across the app
document_analyzer = DocumentAnalyzer()
